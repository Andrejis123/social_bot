"""Regenerate Documents/Social_Bot_Overview.docx and Social_Bot_Technical.docx.

Run from repo root:
    uv run python scripts/_gen_docs.py
"""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.shared import Pt

DOCS_DIR = Path(__file__).resolve().parent.parent / "Documents"


def _style(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)


def _h1(doc: Document, text: str) -> None:
    doc.add_heading(text, level=1)


def _h2(doc: Document, text: str) -> None:
    doc.add_heading(text, level=2)


def _p(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def _bullet(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Bullet")


def _num(doc: Document, text: str) -> None:
    doc.add_paragraph(text, style="List Number")


def _code(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = "Menlo"
    run.font.size = Pt(10)


# ---------------------------------------------------------------------------
# Overview document — for non-engineer audience
# ---------------------------------------------------------------------------

def build_overview() -> Document:
    doc = Document()
    _style(doc)

    doc.add_heading("Social_Bot — Overview", level=0)
    _p(doc,
       "Social_Bot is an automated pipeline that tracks Instagram performance "
       "for client accounts, classifies and describes posts with AI, archives "
       "media in cloud storage, and produces monthly client-facing PowerPoint "
       "reports. It replaces a slower n8n setup with a code-driven workflow "
       "that runs unattended on a VPS.")

    _h1(doc, "What it does")
    _bullet(doc, "Scrapes Instagram posts (weekly) and stories (nightly) for every active client account.")
    _bullet(doc, "Stores engagement metrics as a time series so performance-over-time is queryable.")
    _bullet(doc, "Backs up all media (carousels, Reels covers, story slides) to Supabase Storage.")
    _bullet(doc, "Uses Gemini to classify each post and write a one-paragraph description.")
    _bullet(doc, "Renders a monthly .pptx report per client and publishes it to Supabase + Google Drive.")
    _bullet(doc, "Sends Telegram notifications when runs finish or fail.")

    _h1(doc, "Where things live")
    _h2(doc, "Mac (development)")
    _p(doc, "Code is edited locally on the Mac and pushed to GitHub. No production "
            "workloads run here — only ad-hoc tests, spike scripts, and report previews.")
    _h2(doc, "GitHub")
    _p(doc, "Repository: github.com/Andrejis123/social_bot. The main branch is the "
            "source of truth for what the VPS runs after each `git pull`.")
    _h2(doc, "VPS (DigitalOcean, 161.35.170.254)")
    _p(doc, "All scheduled work runs here. The repo is checked out at /opt/social-bot. "
            "Docker builds the runtime image; root's crontab triggers the jobs.")
    _h2(doc, "Supabase")
    _p(doc, "Postgres database (clients, accounts, posts, stories, metrics, AI fields, "
            "run history) and private Storage bucket (`media`) for downloaded files "
            "and rendered reports.")
    _h2(doc, "Google Drive")
    _p(doc, "Per-client folders receive a copy of each monthly report and a content "
            "bundle (zip of all media for the period) for handoff.")
    _h2(doc, "Telegram")
    _p(doc, "Bot pushes notifications on every cron run — success, partial, or failure — "
            "tagged with the @handle being processed.")

    _h1(doc, "External tools used")
    _bullet(doc, "HikerAPI — primary Instagram scraper (authenticated, sees restricted profiles).")
    _bullet(doc, "Apify — fallback Instagram scraper for accounts Hiker can't reach.")
    _bullet(doc, "Gemini 2.x — AI classification, descriptions, and report narrative synthesis.")
    _bullet(doc, "OpenAI — secondary AI provider, per-client override.")
    _bullet(doc, "Google Drive API — report + content bundle delivery.")
    _bullet(doc, "Telegram Bot API — run notifications.")

    _h1(doc, "Pipeline at a glance")
    _num(doc, "Scrape — pull the latest posts or stories from Instagram (Hiker first, Apify on fallback).")
    _num(doc, "Dedupe — match against the database by (platform, platform_post_id); existing items only get a new metrics snapshot.")
    _num(doc, "Archive — download media files and upload them to Supabase Storage under a per-client/handle/date path.")
    _num(doc, "Classify — Gemini picks a category from the client's `categories.yaml` and writes a short description.")
    _num(doc, "Snapshot — append a row to post_metrics so engagement-over-time is preserved.")
    _num(doc, "Report (monthly) — aggregate the period, synthesize narrative per category, render a .pptx, publish.")
    _num(doc, "Notify — Telegram message summarizing run status and counts.")

    _h1(doc, "What runs when")
    _bullet(doc, "Posts cron: weekly (each active client, all accounts).")
    _bullet(doc, "Stories cron: nightly at 23:00 UTC (restored 2026-06-01).")
    _bullet(doc, "Describe-posts cron: chases unclassified posts after each scrape.")
    _bullet(doc, "Monthly report cron: wired but only triggered manually until the first paying client.")

    _h1(doc, "Repository layout (top level)")
    _code(doc,
          "src/social_bot/    library code (pipeline, scrapers, AI, reports)\n"
          "scripts/           thin CLI entrypoints — what cron calls\n"
          "config/clients/    one folder per client: client.yaml, prompt.md, categories.yaml\n"
          "assets/clients/    per-client brand assets (logo, color palette) — _default/ is the fallback\n"
          "migrations/        SQL schema files applied via Supabase SQL editor\n"
          "docker/            Dockerfile + docker-compose.yml for the VPS runtime\n"
          "Documents/         this overview and the technical companion\n"
          "Generated Reports/ local renders for review before publishing")

    _h1(doc, "Adding a new client (high level)")
    _num(doc, "Copy an existing folder under config/clients/ and edit client.yaml, prompt.md, categories.yaml.")
    _num(doc, "Insert clients + accounts rows in Supabase (matching the slug + handles).")
    _num(doc, "Optionally drop a brand.yaml + logo under assets/clients/<slug>/ for branded reports.")
    _num(doc, "Run a manual scrape from the Mac to sanity-check the config, then commit + push.")

    _h1(doc, "Adding a new platform (high level)")
    _p(doc, "The scrapers folder is a plugin slot. Implementing the Scraper protocol "
            "in a new module and registering it in scrapers/registry.py is enough — "
            "no schema migrations needed, posts/stories tables already carry a platform "
            "column.")

    return doc


# ---------------------------------------------------------------------------
# Technical document — for the engineer maintaining the system
# ---------------------------------------------------------------------------

def build_technical() -> Document:
    doc = Document()
    _style(doc)

    doc.add_heading("Social_Bot — Technical Reference", level=0)
    _p(doc, "Engineer-facing companion to the Overview. Covers the package layout, "
            "the entrypoint scripts cron drives, the end-to-end program flow, and "
            "the VPS runtime in more depth. Current as of 2026-06-13.")

    _h1(doc, "Stack summary")
    _bullet(doc, "Python 3.12, dependency-managed by uv.")
    _bullet(doc, "Typed config via pydantic-settings (config.py) + per-client YAML (config/clients/<slug>/).")
    _bullet(doc, "Supabase (Postgres + Storage) via supabase-py.")
    _bullet(doc, "HikerAPI primary Instagram scraper; Apify (`apify/instagram-scraper`) as fallback.")
    _bullet(doc, "Gemini 2.x (google-genai) primary AI; OpenAI as override.")
    _bullet(doc, "python-pptx for report rendering; google-api-python-client for Drive upload.")
    _bullet(doc, "structlog for JSON logs to stdout.")
    _bullet(doc, "Docker (`docker compose`) wraps the runtime; cron triggers it on the VPS.")

    _h1(doc, "Package layout — src/social_bot/")
    _code(doc,
          "config.py            typed env settings (SUPABASE_*, GEMINI_API_KEY, …)\n"
          "logging.py           structlog setup, injects run_id\n"
          "clients.py           loads config/clients/<slug>/client.yaml into typed objects\n"
          "drive.py             Google Drive uploader (refresh-token OAuth)\n"
          "db/\n"
          "  client.py          cached Supabase client factory\n"
          "  queries.py         typed CRUD — upsert_post, append_metric, …\n"
          "storage/\n"
          "  media.py           download → Supabase Storage; build_storage_path()\n"
          "  reports.py         upload rendered .pptx; signed-URL helpers\n"
          "scrapers/\n"
          "  base.py            Scraper protocol + ScrapedPost / ScrapedMedia dataclasses\n"
          "  _hiker_client.py   HikerAPI HTTP wrapper\n"
          "  instagram.py       multi-tier orchestrator: Hiker → Apify cookie → Apify anon\n"
          "  registry.py        platform → Scraper mapping\n"
          "ai/\n"
          "  classifier.py      classify(post, media_sample, client_cfg) → {category, …}\n"
          "  descriptor.py      one-paragraph post/story description\n"
          "  media_sampler.py   pick_for_ai() — which slides to send (carousels, reels, single)\n"
          "  providers/\n"
          "    gemini.py        google-genai wrapper (thinking_budget=0 for prose tasks)\n"
          "    openai.py        OpenAI fallback\n"
          "pipeline/\n"
          "  run_context.py     opens/closes run_history row, collects per-item errors\n"
          "  ingest_posts.py    weekly job core: scrape → dedupe → media → AI → metric\n"
          "  ingest_stories.py  nightly job core: same shape, story tables\n"
          "  describe_posts.py  generates AI descriptions for classified posts\n"
          "  describe_stories.py same for stories\n"
          "reports/\n"
          "  brand.py           loads assets/clients/<slug>/brand.yaml\n"
          "  data.py            build_period(start, end), load_report_data(client, period)\n"
          "  synthesis.py       two-pass Gemini synthesis: cluster items → narrative + image pick\n"
          "  theme.py           geometric constants (inches) — single source of truth for slide geometry\n"
          "  layouts.py         per-layout slide builders\n"
          "  renderer.py        publish_report(client, period) — render + upload + Drive + Telegram\n"
          "notifications/\n"
          "  telegram.py        notify_run_started / _completed / _failed / _ai_exhausted")

    _h1(doc, "Entrypoint scripts (what cron calls)")
    _h2(doc, "scripts/scrape_posts.py")
    _p(doc, "Thin CLI wrapper around pipeline.ingest_posts. Args: --client <slug>, "
            "optional --account, --since, --until, --limit. Iterates the client's "
            "active accounts and runs ingest_posts_for_client() per account.")
    _h2(doc, "scripts/scrape_stories.py")
    _p(doc, "Same shape as scrape_posts but for ephemeral story content. Stories "
            "have separate tables (stories, story_media) and a different scraper "
            "tier (cookie-based on Apify in the fallback path).")
    _h2(doc, "scripts/describe_posts.py / describe_stories.py")
    _p(doc, "Walk recently-classified rows that still lack ai_description and call "
            "the descriptor pipeline. --sleep N throttles between calls to stay "
            "under Gemini quotas. Notifications are per-account, labeled \"AI Descriptions\".")
    _h2(doc, "scripts/run_monthly_reports.py")
    _p(doc, "Monthly report driver. Usage: "
            "`python -m scripts.run_monthly_reports <YYYY-MM-DD> <YYYY-MM-DD> [client …]`. "
            "Builds a Period (explicit window, not calendar month — matches the "
            "\"25 April – 25 May\" convention), then calls publish_report() per client. "
            "Default client set is [agape, ecig-monitoring, iluminatecz].")
    _h2(doc, "scripts/make_content_bundle.py")
    _p(doc, "Bundles all scraped media for a client + window into a zip and uploads "
            "to the client's Drive folder. Runs alongside the monthly report.")
    _h2(doc, "scripts/retry_ai.py")
    _p(doc, "Re-runs classification/description for rows that failed the last time "
            "(quota / transient error). Used to drain backlog after an outage.")
    _h2(doc, "scripts/_google_auth.py")
    _p(doc, "One-time helper to mint a Google Drive refresh token from "
            "credentials.json (downloaded from GCP). Writes the token to .env.")
    _h2(doc, "scripts/_backfill_reel_covers.py, _migrate_hiker_dedup.py, backfill_storage_paths.py")
    _p(doc, "One-off maintenance scripts. Reel covers are stored at slide_index=99 "
            "at scrape time (the lazy heal-at-report-time approach was unreliable "
            "because Instagram signed URLs are session-bound).")

    _h1(doc, "Program flow — ingest_posts (weekly)")
    _code(doc,
          "1. open run_history row (status='running', job='ingest_posts', client_slug)\n"
          "2. load_client(slug) → typed ClientConfig\n"
          "3. for each Instagram account on the client:\n"
          "     a. registry.get('instagram').scrape_posts(account, limit, since)\n"
          "        - Hiker tier 1\n"
          "        - Apify cookie tier 2 (if Hiker fails / restricted)\n"
          "        - Apify anon tier 3 (last resort)\n"
          "     b. for each ScrapedPost (each wrapped in try/except → run_item_errors):\n"
          "          - queries.find_post_by_platform_id(platform, platform_post_id)\n"
          "          - if EXISTS: queries.append_post_metric(post_id, snapshot)\n"
          "          - if NEW:\n"
          "              * queries.insert_post(...)  with raw_payload jsonb\n"
          "              * for each ScrapedMedia: storage.media.download_and_upload(); insert media\n"
          "              * (reels) capture cover at slide_index=99 alongside the video\n"
          "              * ai.media_sampler.pick_for_ai(post) → subset\n"
          "              * ai.classifier.classify(post, sample, client_cfg)\n"
          "              * queries.update_post_ai(post_id, category, prompt_version)\n"
          "              * queries.append_post_metric(post_id, first snapshot)\n"
          "4. close run_history (success / partial / failed) + Telegram notify")

    _h1(doc, "Program flow — monthly report")
    _code(doc,
          "1. build_period(start, end) → Period(start, end, label)\n"
          "2. for each client in arg list:\n"
          "     a. data.load_report_data(client, period) → posts, metrics, brand, top items\n"
          "     b. synthesis.synthesize_category(items) — two-pass:\n"
          "          pass 1: cluster posts into report items (Gemini)\n"
          "          pass 2: per-item narrative + image pick (Gemini, coupled)\n"
          "     c. renderer.publish_report(...):\n"
          "          - layouts build slides from theme.py geometry constants\n"
          "          - python-pptx writes /tmp/reports/<slug>_<period>.pptx\n"
          "          - storage.reports uploads to Supabase, returns signed URL\n"
          "          - drive.upload mirrors into client's Drive folder (best-effort)\n"
          "          - notifications.telegram.notify_report_published")
    _p(doc, "Style rules enforced in synthesis prompts and rendering: no em-dashes "
            "or dash-substitutes anywhere; the report subject always renders as the "
            "verbatim @handle (never the client name/slug).")

    _h1(doc, "Database schema highlights")
    _bullet(doc, "clients(id, slug, name, is_active)")
    _bullet(doc, "accounts(id, client_id, platform, handle, platform_account_id, is_owned, is_active) UNIQUE(platform, handle)")
    _bullet(doc, "posts(id, account_id, platform, platform_post_id, post_type, caption, permalink, posted_at, first_seen_at, raw_payload jsonb, ai_category, ai_description, ai_analyzed_at, ai_prompt_version) UNIQUE(platform, platform_post_id)")
    _bullet(doc, "post_metrics(post_id, scraped_at, like_count, comment_count, view_count, play_count, save_count) — append-only time series, indexed (post_id, scraped_at desc)")
    _bullet(doc, "media(post_id, slide_index, media_type, source_url, storage_path, …) UNIQUE(post_id, slide_index); slide_index=99 reserved for reel covers")
    _bullet(doc, "stories + story_media — same shape as posts/media; separate because stories are ephemeral with different fields")
    _bullet(doc, "run_history(job_name, client_slug, started_at, finished_at, status, items_new/updated/failed, error_summary)")
    _bullet(doc, "run_item_errors(run_id, item_ref, stage, error_message) — one post failing never kills the run")
    _p(doc, "Migrations live in migrations/: 0001_initial_schema.sql, "
            "0002_add_ai_description.sql, 0003_add_stories_ai.sql. Applied "
            "manually via the Supabase SQL editor.")

    _h1(doc, "Storage path scheme")
    _code(doc, "media/{client_slug}/{platform}/{account_handle}/{YYYY}/{MM}/{post_id}/{slide_index}.{ext}")
    _p(doc, "Sortable, human-browsable, and prefix-queryable. Reports are stored "
            "separately under reports/{client_slug}/{period}.pptx.")

    _h1(doc, "VPS runtime (DigitalOcean, 161.35.170.254)")
    _h2(doc, "Filesystem")
    _bullet(doc, "/opt/social-bot — git checkout of main")
    _bullet(doc, "/opt/social-bot/.env — production secrets (not in git)")
    _bullet(doc, "/opt/social-bot/crontab.backup.* — versioned backups of root's crontab")
    _h2(doc, "Docker")
    _bullet(doc, "Image tag: social-bot:latest, built from docker/Dockerfile")
    _bullet(doc, "Compose file: docker/docker-compose.yml — single-shot `app` service, env_file ../.env")
    _bullet(doc, "Build after pull: `docker compose -f docker/docker-compose.yml build`")
    _bullet(doc, "Run a job: `docker compose -f docker/docker-compose.yml run --rm app python -m scripts.scrape_posts --client <slug>`")
    _h2(doc, "Cron")
    _bullet(doc, "Owned by root. Edit with `crontab -e`; always back up first to /opt/social-bot/crontab.backup.<ISO>.")
    _bullet(doc, "Posts cron: weekly per active client (all accounts).")
    _bullet(doc, "Stories cron: nightly 23:00 UTC (restored 2026-06-01).")
    _bullet(doc, "Describe cron: follows posts/stories to drain unclassified rows.")
    _bullet(doc, "Monthly report cron: wired, currently disabled until the first paying client.")
    _h2(doc, "Deploy flow")
    _code(doc,
          "# Mac\n"
          "git push origin main\n\n"
          "# VPS (root@161.35.170.254:/opt/social-bot)\n"
          "git pull\n"
          "docker compose -f docker/docker-compose.yml build\n"
          "# cron picks up the new image on the next scheduled run")

    _h1(doc, "Operational notes & gotchas")
    _bullet(doc, "Apify pricing: headline \"$X/1000\" can hide per-event surcharges. One careless run can blow the Free tier's $5/month cap and block all subsequent runs until reset — always verify the pricing model before enabling a new actor.")
    _bullet(doc, "Gemini thinking tokens: pass thinking_budget=0 for classify/describe/synthesis prose tasks. Thinking tokens count against max_output_tokens AND are billed.")
    _bullet(doc, "Instagram signed URLs are session-bound. Reel covers must be captured at scrape time (slide_index=99), not lazily at report time.")
    _bullet(doc, "Telegram notifications: always tag with the @handle, never the client slug or name.")
    _bullet(doc, "Report prose: no em-dashes anywhere; bump prompt_version when changing handle/style rules so re-runs are targetable.")

    _h1(doc, "Testing & verification")
    _p(doc, "tests/ contains unit tests for normalizers (Hiker, Instagram, Apify) and "
            "the media sampler. Run with `uv run python -m pytest tests/ -q`. "
            "Larger end-to-end testing strategy is the next planning block "
            "(decided 2026-06-12) — scope discipline, /security-review, /ultrareview "
            "to be brainstormed before implementation.")

    return doc


def main() -> None:
    DOCS_DIR.mkdir(parents=True, exist_ok=True)
    overview = build_overview()
    overview_path = DOCS_DIR / "Social_Bot_Overview.docx"
    overview.save(overview_path)

    technical = build_technical()
    technical_path = DOCS_DIR / "Social_Bot_Technical.docx"
    technical.save(technical_path)

    print(f"wrote {overview_path}")
    print(f"wrote {technical_path}")


if __name__ == "__main__":
    main()
